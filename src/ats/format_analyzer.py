"""
ATS Format Analyzer - Detects ATS-unfriendly formatting elements.
Analyzes document structure for compatibility issues.
"""

import re
from typing import Dict, List, Tuple
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

# Optional dependencies for advanced analysis
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not available. Limited PDF analysis.")

try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Limited DOCX analysis.")


class FormatAnalyzer:
    """
    Analyzes resume format for ATS compatibility.
    Detects problematic elements like tables, images, and complex formatting.
    """
    
    def __init__(self):
        """Initialize format analyzer."""
        self.supported_formats = {'.pdf', '.docx', '.doc', '.txt'}
        self.preferred_formats = {'.pdf', '.docx'}
    
    def analyze_file(self, file_path: str) -> Dict:
        """
        Analyze file format for ATS compatibility.
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Dict with format analysis results
        """
        path = Path(file_path)
        
        if not path.exists():
            return self._error_result(f"File not found: {file_path}")
        
        file_extension = path.suffix.lower()
        
        # File format assessment
        format_issues = []
        format_score = 100
        
        if file_extension not in self.supported_formats:
            format_issues.append({
                'severity': 'critical',
                'issue': f'Unsupported file format: {file_extension}',
                'impact': 'ATS systems may not be able to parse this file',
                'recommendation': 'Convert to PDF or DOCX format'
            })
            format_score = 0
        elif file_extension not in self.preferred_formats:
            format_issues.append({
                'severity': 'medium',
                'issue': f'Non-preferred format: {file_extension}',
                'impact': 'Some ATS systems have limited support',
                'recommendation': 'Consider using PDF or DOCX for better compatibility'
            })
            format_score = 70
        
        # Analyze document structure based on format
        if file_extension == '.pdf' and PDF_AVAILABLE:
            structure_analysis = self._analyze_pdf_structure(file_path)
        elif file_extension in ['.docx', '.doc'] and DOCX_AVAILABLE:
            structure_analysis = self._analyze_docx_structure(file_path)
        else:
            structure_analysis = self._basic_structure_analysis()
        
        # Combine results
        all_issues = format_issues + structure_analysis['issues']
        
        # Calculate final format score
        structure_penalty = (100 - structure_analysis['score']) * 0.7
        final_score = max(0, format_score - structure_penalty)
        
        return {
            'file_format': file_extension,
            'format_score': round(final_score, 2),
            'is_ats_friendly': final_score >= 70,
            'issues': all_issues,
            'structure_details': structure_analysis['details'],
            'recommendations': self._generate_format_recommendations(all_issues)
        }
    
    def _analyze_pdf_structure(self, file_path: str) -> Dict:
        """Analyze PDF document structure."""
        issues = []
        details = {}
        score = 100
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                details['num_pages'] = num_pages
                details['has_images'] = False
                details['has_forms'] = False
                
                # Check page count
                if num_pages > 2:
                    issues.append({
                        'severity': 'medium',
                        'issue': f'Resume is {num_pages} pages long',
                        'impact': 'ATS systems may truncate or score lower',
                        'recommendation': 'Keep resume to 1-2 pages for optimal ATS performance'
                    })
                    score -= 10
                
                # Check for images (basic heuristic)
                for page_num, page in enumerate(pdf_reader.pages):
                    if '/XObject' in page.get('/Resources', {}):
                        xobjects = page['/Resources']['/XObject'].get_object()
                        for obj in xobjects:
                            if xobjects[obj]['/Subtype'] == '/Image':
                                details['has_images'] = True
                                issues.append({
                                    'severity': 'high',
                                    'issue': 'Document contains images',
                                    'impact': 'Images are not readable by ATS systems',
                                    'recommendation': 'Remove images, photos, logos, and graphics'
                                })
                                score -= 20
                                break
                        if details['has_images']:
                            break
                
                # Check for form fields
                if '/AcroForm' in pdf_reader.trailer['/Root']:
                    details['has_forms'] = True
                    issues.append({
                        'severity': 'medium',
                        'issue': 'Document contains form fields',
                        'impact': 'Form fields may not be parsed correctly',
                        'recommendation': 'Flatten form fields or use plain text'
                    })
                    score -= 15
                
                # Check for encryption
                if pdf_reader.is_encrypted:
                    issues.append({
                        'severity': 'critical',
                        'issue': 'PDF is encrypted/password-protected',
                        'impact': 'ATS cannot parse encrypted documents',
                        'recommendation': 'Remove password protection'
                    })
                    score = 0
                
        except Exception as e:
            logger.error(f"Error analyzing PDF structure: {e}")
            issues.append({
                'severity': 'high',
                'issue': 'PDF analysis failed',
                'impact': 'Document may have structural issues',
                'recommendation': 'Ensure PDF is not corrupted'
            })
            score = 50
        
        return {
            'score': max(0, score),
            'issues': issues,
            'details': details
        }
    
    def _analyze_docx_structure(self, file_path: str) -> Dict:
        """Analyze DOCX document structure."""
        issues = []
        details = {}
        score = 100
        
        try:
            doc = Document(file_path)
            
            # Count structural elements
            num_tables = len(doc.tables)
            num_paragraphs = len(doc.paragraphs)
            num_sections = len(doc.sections)
            
            details['num_tables'] = num_tables
            details['num_paragraphs'] = num_paragraphs
            details['num_sections'] = num_sections
            details['has_headers'] = False
            details['has_footers'] = False
            details['has_columns'] = False
            
            # Check for tables
            if num_tables > 0:
                issues.append({
                    'severity': 'high',
                    'issue': f'Document contains {num_tables} table(s)',
                    'impact': 'Tables often cause parsing errors in ATS systems',
                    'recommendation': 'Convert tables to simple text with clear formatting'
                })
                score -= 25
            
            # Check for headers/footers
            for section in doc.sections:
                if section.header.paragraphs and any(p.text.strip() for p in section.header.paragraphs):
                    details['has_headers'] = True
                if section.footer.paragraphs and any(p.text.strip() for p in section.footer.paragraphs):
                    details['has_footers'] = True
            
            if details['has_headers'] or details['has_footers']:
                issues.append({
                    'severity': 'medium',
                    'issue': 'Document contains headers or footers',
                    'impact': 'Content in headers/footers may be ignored by ATS',
                    'recommendation': 'Move important information to document body'
                })
                score -= 15
            
            # Check for multi-column layout
            for section in doc.sections:
                if hasattr(section, '_sectPr') and section._sectPr.xpath('.//w:cols[@w:num]'):
                    col_elements = section._sectPr.xpath('.//w:cols[@w:num]')
                    if col_elements and int(col_elements[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num', '1')) > 1:
                        details['has_columns'] = True
                        break
            
            if details['has_columns']:
                issues.append({
                    'severity': 'high',
                    'issue': 'Document uses multi-column layout',
                    'impact': 'Column layouts often cause content to be read out of order',
                    'recommendation': 'Use single-column layout for ATS compatibility'
                })
                score -= 20
            
            # Check for embedded objects (images, charts)
            has_images = False
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    has_images = True
                    break
            
            if has_images:
                details['has_images'] = True
                issues.append({
                    'severity': 'high',
                    'issue': 'Document contains images or graphics',
                    'impact': 'Visual elements are not readable by ATS',
                    'recommendation': 'Remove all images, charts, and graphics'
                })
                score -= 20
            
        except Exception as e:
            logger.error(f"Error analyzing DOCX structure: {e}")
            issues.append({
                'severity': 'medium',
                'issue': 'DOCX analysis encountered errors',
                'impact': 'Document may have compatibility issues',
                'recommendation': 'Ensure document is not corrupted'
            })
            score = 60
        
        return {
            'score': max(0, score),
            'issues': issues,
            'details': details
        }
    
    def _basic_structure_analysis(self) -> Dict:
        """Basic structure analysis when libraries unavailable."""
        return {
            'score': 70,
            'issues': [{
                'severity': 'low',
                'issue': 'Limited format analysis available',
                'impact': 'Unable to detect all formatting issues',
                'recommendation': 'Install PyPDF2 and python-docx for comprehensive analysis'
            }],
            'details': {'analysis_mode': 'basic'}
        }
    
    def _generate_format_recommendations(self, issues: List[Dict]) -> List[str]:
        """Generate actionable format recommendations."""
        recommendations = []
        
        # Categorize issues
        has_critical = any(i['severity'] == 'critical' for i in issues)
        has_high = any(i['severity'] == 'high' for i in issues)
        has_tables = any('table' in i['issue'].lower() for i in issues)
        has_images = any('image' in i['issue'].lower() for i in issues)
        has_columns = any('column' in i['issue'].lower() for i in issues)
        
        if has_critical:
            recommendations.append("Address critical issues immediately - document may not be parseable")
        
        if has_tables:
            recommendations.append("Replace tables with simple text formatting using bullets and indentation")
        
        if has_images:
            recommendations.append("Remove all visual elements and replace with text descriptions where necessary")
        
        if has_columns:
            recommendations.append("Convert to single-column layout for proper content sequencing")
        
        if not issues:
            recommendations.append("Format is ATS-friendly - no structural changes needed")
        
        return recommendations
    
    def _error_result(self, error_message: str) -> Dict:
        """Return error result."""
        return {
            'file_format': 'unknown',
            'format_score': 0,
            'is_ats_friendly': False,
            'issues': [{
                'severity': 'critical',
                'issue': error_message,
                'impact': 'Cannot analyze file',
                'recommendation': 'Verify file path and format'
            }],
            'structure_details': {},
            'recommendations': ['Fix file access issues before proceeding']
        }
    
    def analyze_text_extractability(self, extracted_text: str) -> Dict:
        """
        Analyze how well text was extracted (quality check).
        
        Args:
            extracted_text: Text extracted from resume
            
        Returns:
            Dict with extractability analysis
        """
        issues = []
        score = 100
        
        if not extracted_text or len(extracted_text.strip()) < 50:
            issues.append({
                'severity': 'critical',
                'issue': 'Very little text was extracted',
                'impact': 'ATS will have minimal content to analyze',
                'recommendation': 'Ensure resume uses standard fonts and text (not images of text)'
            })
            score = 0
        else:
            # Check for excessive special characters
            special_char_ratio = len(re.findall(r'[^\w\s.,;:()\-/]', extracted_text)) / len(extracted_text)
            if special_char_ratio > 0.05:
                issues.append({
                    'severity': 'medium',
                    'issue': f'High special character density ({special_char_ratio*100:.1f}%)',
                    'impact': 'May indicate encoding or formatting issues',
                    'recommendation': 'Use standard characters and avoid decorative fonts'
                })
                score -= 20
            
            # Check for unusual whitespace
            lines = extracted_text.split('\n')
            empty_line_ratio = sum(1 for line in lines if not line.strip()) / max(len(lines), 1)
            if empty_line_ratio > 0.5:
                issues.append({
                    'severity': 'low',
                    'issue': 'Excessive blank lines detected',
                    'impact': 'May indicate parsing issues with document structure',
                    'recommendation': 'Reduce excessive spacing between sections'
                })
                score -= 10
            
            # Check for broken words (words with spaces in middle)
            broken_words = re.findall(r'\b[a-z]+ [a-z]+(?=[A-Z]|\s|$)', extracted_text)
            if len(broken_words) > 5:
                issues.append({
                    'severity': 'medium',
                    'issue': 'Detected word spacing issues',
                    'impact': 'Text may not be extracted correctly',
                    'recommendation': 'Avoid complex formatting that breaks word boundaries'
                })
                score -= 15
        
        return {
            'extractability_score': max(0, score),
            'text_length': len(extracted_text),
            'is_extractable': score >= 50,
            'issues': issues
        }