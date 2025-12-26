"""
DOCX parser implementation
"""

from pathlib import Path
from typing import Optional
from docx import Document

from src.parser.base_parser import BaseParser
from src.utils.text_utils import clean_text


class DOCXParser(BaseParser):
    """Parser for DOCX (Microsoft Word) files."""
    
    def supports_file_type(self, file_path: Path) -> bool:
        """Check if file is DOCX."""
        return file_path.suffix.lower() in ['.docx', '.doc']
    
    def parse(self, file_path: Path) -> Optional[str]:
        """
        Parse DOCX file and extract text.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text or None if parsing fails
        """
        if not self.validate_file(file_path):
            return None
        
        if not self.supports_file_type(file_path):
            self.logger.error(f"File is not a DOCX: {file_path}")
            return None
        
        try:
            text = self._extract_text(file_path)
            
            if text:
                self.logger.info(f"Successfully parsed DOCX: {file_path.name}")
                return clean_text(text)
            else:
                self.logger.warning(f"No text extracted from DOCX: {file_path}")
                return None
                
        except Exception as e:
            self.logger.error(f"Failed to parse DOCX {file_path}: {e}")
            return None
    
    def _extract_text(self, file_path: Path) -> Optional[str]:
        """
        Extract text from DOCX preserving structure.
        
        Args:
            file_path: Path to DOCX
            
        Returns:
            Extracted text
        """
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            text = '\n'.join(text_parts)
            self.logger.debug(f"Extracted {len(text)} characters from DOCX")
            
            return text if text else None
            
        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {e}")
            return None
    
    def get_paragraph_count(self, file_path: Path) -> Optional[int]:
        """
        Get number of paragraphs in DOCX.
        
        Args:
            file_path: Path to DOCX
            
        Returns:
            Paragraph count or None if error
        """
        try:
            doc = Document(file_path)
            return len([p for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            self.logger.warning(f"Could not get paragraph count: {e}")
            return None